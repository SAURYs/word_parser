"""
将.docx文件转为pdf后进行解析、以及使用python-docx工具包进行解析
解析的具体功能有：
1.页码ID格式正确性
2.表格标题及表格内容格式检查
3.图片标题格式正确性
4.公式格式正确性
5.注释格式正确性
6.附录编号 附录格式正确性
7.参考文献格式正确性


"""
import os
from io import open
from string import digits
# from numpy import shape, argmax
from docx import Document
from docx.shared import Cm, Pt,Emu
import re

from docx.shared import Pt
from docx.shared import Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

normal_font_name = None
normal_Chinese_font_name = None

class get_doc_default_Pr:
    '''
    获取docx文件默认参数的类。
    默认的参数主要有：
                rPr: font_name, font_size, Chinese_font_name;
                pPr；
    '''
    @classmethod
    def font_name(cls,path):
        '''
        :param path: 文件的绝对路径
        :return: rPr->rFonts.ascii
        '''
        doc = Document(path)
        element= doc._part._styles_part._element
        for child in element.iter():
            if child.tag.endswith('docDefaults'):
                for i in child.iter():
                    #print(i.tag)
                    if i.tag.endswith('rFonts'):
                        if hasattr(i,'ascii'):
                            return i.ascii
                        else:
                            return None

    @classmethod
    def Chinese_font_name(cls, path):
        '''
                :param path: 文件的绝对路径
                :return: rPr->rFonts.eastAsia
                '''
        doc = Document(path)
        element = doc._part._styles_part._element
        for child in element.iter():
            if child.tag.endswith('docDefaults'):
                for i in child.iter():
                    #print(i.tag)
                    if i.tag.endswith('rFonts'):
                        if hasattr(i,'eastAsia'):
                            return i.eastAsia
                        else:
                            return None

def show_all_content_fmt(path):
    '''
    输出所有段落格式便于调试
    '''
    for i in Document(path).paragraphs:
        print(get_paragraph_format(i))


def is_all_zh(string):
    '''
    判断是否全为中文字符
    :param s: string 字符串
    :return: bool
    '''
    for c in string:
        if not ('\u4e00' <= c <= '\u9fa5'):
            return False
    return True

def extract_pageid(page_content, mode=0):
    '''
    # 获取页码
    input:page_content为一页的内容列表
          mode为页码位置，默认为0说明页码在最后一行（页脚），为1说明在页眉
    output:str格式，该页页码
    但页码是在左边还是右边无法判断！
    '''
    if len(page_content) == 0:
        return '-1'
    elif not mode:
        return get_text_page(page_content[-1])
    else:
        return get_text_page(page_content[0])


def get_text_page(text):
    '''
    # 获取文档内容
    # 这是基于pdf处理的部分
    :param text: 提取出的一页的内容，一行为一个列表
    :return: 返回的是str类型，若页码不符合规范，返回'-1'，正确返回其ID
    '''
    pageID = ''
    if (text[-1].isdigit()==False):
        text_number = '-1'
    elif len(text) > 5:
        text_rev = list(reversed(text))
        for i in range(len(text_rev)):
            if text_rev[i].isdigit():
                pageID += text_rev[i]
                if not text_rev[i + 1].isdigit():
                    break
        pageID = pageID[::-1]
    else:
        pageID = '-1'
    return pageID


def get_table_format(path):
    '''
    # 使用python-docx判断表格内容的格式正确性
    :param path: docx的文件的绝对路径
    :return: table_font_name, table_font_size 字体、字号(Pt)
    '''
    fmt = []
    document = Document(path)
    tables = document.tables
    table_nums = len(tables)
    # 遍历每一个表格
    fmt = [[] for i in range(table_nums)]
    for table_index,table_content in enumerate(tables):
        # 遍历每个表格中的内容
        for row in table_content.rows:
            for cell in row.cells:
                _paragraphs = cell.paragraphs
                for each_paragraphs in _paragraphs:
                    fmt[table_index].append(get_paragraph_format(each_paragraphs))
    return fmt


def get_tableTitle_format(paragraph):
    '''
    通过传入Document.paragraphs中的一个段落，来获得表的内容、字体名称、字体大小、段前距离、段后距离、首行缩进、是否居中等信息
    图表还要检查是否序号与表名之间空一字,ie. 表2-2 物种的起源 不可以是：表2-2物种的起源,或者多空格的情况
    但现在还有一行含有多个图表的情况[判断连续性的时候要分开才行]
    :param paragraph:  Document.paragraphs中的一个段落
    :return
    '''
    each_paragraph = paragraph
    paragraph_content = each_paragraph.text
    content_without_space = paragraph_content.replace(' ', '')
    table_format = []
    pattern = '^[表]\d+[--]\d+'

    if content_without_space=='':
        pass
    elif re.match(pattern, content_without_space) and len(content_without_space) < 50:
        table_format = get_paragraph_format(each_paragraph)
    return table_format


def get_imageTitle_format(paragraph):
    '''
    通过传入Document.paragraphs中的一个段落，来获得图标题的内容、字体名称、字体大小、段前距离、段后距离、首行缩进、是否居中等信息
    图表还要检查是否序号与表名之间空一字,ie. 图2-2 物种的起源 不可以是：图2-2物种的起源,或者多空格的情况
    但现在还有一行含有多个图表的情况[判断连续性的时候要分开才行]
    :param paragraph:  Document.paragraphs中的一个段落
    :return
    '''
    each_paragraph = paragraph
    paragraph_content = each_paragraph.text
    image_format = []
    content_without_space = paragraph_content.replace(' ', '')
    pattern = '^[图]\d+[--]\d+'
    if content_without_space=='':
        pass
    elif re.match(pattern, content_without_space) and len(content_without_space) < 50:
        image_format = get_paragraph_format(each_paragraph)

    return image_format


def get_xubiao_format(paragraph):
    '''
    获取续表标题及格式
    :param paragraph:  Document.paragraphs中的一个段落
    :return
    '''
    xubiao_format = []
    each_paragraph = paragraph
    paragraph_content = each_paragraph.text
    content_without_space = paragraph_content.replace(' ', '')
    pattern = '^续表\d+[--]\d+'
    if content_without_space=='':
        pass
    elif re.match(pattern, content_without_space) and len(content_without_space) < 50:
        xubiao_format = get_paragraph_format(each_paragraph)
    return xubiao_format


def table_title_has_space(title):
    '''
    该函数是判断表2-2 物种的起源这种，序号与图、表名称之间是否空1字
    '''
    pattern = '\d+[--]\d+'
    search_res = re.search(pattern, title)
    if search_res == None:
        return False
    else:
        index = search_res.end()
    if title[index]==' ':
        return True
    else:
        return False

def get_header_and_footer(path):
    '''

    :param path: docx文件的绝对路径
    :return: 页眉 和 页脚
    注：自动生成的页码虽然是包含在页眉或者页脚中，但通过获取页眉或者页脚的文字无法得到页码。
    '''
    font_name=[]
    font_size=[]
    footer_content=[]
    document = Document(path)
    sections = document.sections
    section_nums = len(sections)
    print('section:', end='')
    print(section_nums)
    for each_section in sections:
        header = each_section.header
        footer = each_section.footer
        print('footer',end =':')
        print(footer)

        #header_paragraphs = header.paragraphs
        footer_paragraphs = footer.paragraphs
        for i in footer_paragraphs:
            runs = i.runs
            for each_run in runs:
                font_name.append(each_run.font.name)
                font_size.append(each_run.font.size/12700)
                footer_content.append(each_run.text)
    return font_name, font_size,footer_content

def get_appendix_format(path):
    '''
    !!!附录 二字 是宋体 目前检测为None
    secondary_appendix_index_list:二级附录诸如附录A，附录B等的索引列表
    secondary_appendix_format:二级附录的格式规范列表
    appendix_format: 附录 二字的格式规范
    secondary_appendix_num： 二级附录的个数
    :param path: 文件路径
    :return: appendix_format, secondary_appendix_format（列表）
    '''
    doc = Document(path)
    pars = doc.paragraphs
    secondary_appendix_index_list = []
    secondary_appendix_format = []
    for i in range(len(pars)):
        if re.sub(' ', '', pars[i].text) == '附录': #去掉附录二字中的空格
            appendix_format = get_paragraph_format(pars[i])
        if re.match(r'附录\w{1,2}', pars[i].text.split(' ', 1)[0]):  #匹配二级附录（先分割后匹配）
            secondary_appendix_index_list.append(i)
    secondary_appendix_num = len(secondary_appendix_index_list)
    for j in range(secondary_appendix_num):
        secondary_appendix_format.append(get_paragraph_format(pars[secondary_appendix_index_list[j]]))
    return appendix_format, secondary_appendix_format

def get_reference_format(path):
    '''
    !!!文献内容的字体检测为none
    reference_format：参考文献 四个字的格式规范元组
    reference_content_format：文献内容规范格式列表
    reference_index：参考文献 四个字的索引
    reference_content_index： 文献开始的索引
    :param path:文件路径
    :return:reference_format , reference_content_format
    '''
    doc = Document(path)
    pars = doc.paragraphs
    reference_index = None
    reference_content_format = []
    for index, elem in enumerate(pars):
        if re.sub(' ', '', pars[index].text) == '参考文献':
            reference_index = index
    if reference_index == None:
        return None
    else:
        for i in pars[reference_index:]:
            reference_content_format.append(get_paragraph_format(i))
    return reference_content_format



def get_textbox_content_fontname_font_size(path):
    '''
    # 文本框中也有一般的正常文本，并不全是插图！！该函数是获得文本框中的内容，及其字体、字号。
    :param path: docx文件的绝对路径
    :return: content文本框的内容 ，font_name文字名称，font_size文字大小
    '''
    content = []
    font_size = []
    font_name = []
    document = Document(path)
    for child in document.element.body.iter():
        #print(child)
        if child.tag.endswith('textbox'):
            # c是textbox中的元素
            for c in child.iter():
                c_tag = c.tag
                # 遇到段落换行
                if c_tag.endswith('main}pPr'):
                    pass
                # 在一个run中做操作
                elif c_tag.endswith('main}r'):
                    #content.append(c.text)
                    # 如果text为空则不去判断其字体及字体大小
                    tmp = c.text
                    #除去全为空格的子串
                    if tmp.isspace():
                        tmp = ''
                    if tmp!='':
                        content.append(tmp)
                        # i 是字符段不为空时 run中的元素
                        for i in c.iter():
                            if i.tag.endswith('main}rPr'):
                                if i.rFonts == None:
                                    name ='None'
                                else:
                                    name = i.rFonts.ascii
                                    if name ==None:
                                        name = i.rFonts.hAnsi
                                if i.sz ==None:
                                    size = 'None'
                                else:
                                    size = i.sz.val/12700
                        if name ==None:
                            name ='None'
                        font_name.append(name)
                        font_size.append(size)
    return content, font_name, font_size


def get_paragraph_identation(path):
    '''
    获得段落的段前段后距离(注，这里的段前段后是加上空行之后的段前段后驹，即两个相邻段落之间的距离）
    :param path: docx文件的绝对路径
    要注意判断首行缩进2字符（正文段落，非标题）
    :return:
    '''
    document =Document(path)
    pass

def get_style_format(style):
    '''
    鉴于style存在多次继承的情况，该函数会找到最近一次继承的值，直到找不到继承的样式为止
    :param style: 样式
    :return: 属性值

    '''
    global normal_font_name
    global normal_Chinese_font_name
    if style.name == 'Normal':
        if style.font.name == None:
            style.font.name = normal_font_name
        if style.font.Chinese_font_name == None:
            style.font.Chinese_font_name = normal_Chinese_font_name
    font_name = style.font.name
    Chinese_font_name = style.font.Chinese_font_name
    space_before = style.paragraph_format.space_before
    space_after = style.paragraph_format.space_after
    left_indent = style.paragraph_format.left_indent
    right_indent = style.paragraph_format.right_indent
    first_line_indent = style.paragraph_format.first_line_indent
    alignment = style.paragraph_format.alignment
    while font_name == None or Chinese_font_name == None or space_before == None or \
            space_after == None or left_indent == None or right_indent == None\
            or first_line_indent ==None or alignment == None :
        style = style.base_style
        if style:
            if font_name == None:
                font_name = style.font.name
            if Chinese_font_name == None:
                Chinese_font_name = style.font.Chinese_font_name
            if space_before == None:
                space_before = style.paragraph_format.space_before
            if space_after == None:
                space_after = style.paragraph_format.space_after
            if left_indent == None:
                left_indent = style.paragraph_format.left_indent
            if right_indent == None:
                right_indent = style.paragraph_format.right_indent
            if first_line_indent == None:
                first_line_indent = style.paragraph_format.first_line_indent
            if alignment == None:
                alignment = style.paragraph_format.alignment


        else:
            return [font_name, Chinese_font_name,alignment, space_before, space_after,left_indent,right_indent,first_line_indent]
    return [font_name, Chinese_font_name, space_before, space_after,left_indent,right_indent,first_line_indent]

def get_numbered_title(path):
    '''
    说明：此函数：
            1.获得章节标题的格式
            需要注意的是中文字符中的点即1.1是无法识别出来的
            2. 获得表格标题、插图标题的格式(目前暂不支持文本框的插图标题，同样也不支持穿插在文本段落中的插图标题（要另起一行才行），
                若出现此种情况，一律认为图表标题不存在。[后续可做改进]

    :param path: docx文件的绝对路径
    一级标题、二级标题、三级标题、四级标题的字体、字号、段前段后距离、对齐方式
    :return:
    '''
    document = Document(path)
    _pattern  = r'^\第.\章'
    _pattern0 = r'^\d\.'
    _pattern1 = r'^\d\.\d'
    _pattern2 = r'^\d\.\d\.\d'
    _pattern3 = r'^\d\.\d\.\d\.\d'

    auto_numbered_title_1 = []  # 形如x.x
    auto_numbered_title_2 = []  # 形如x.x.x
    auto_numbered_title_3 = []  # 形如x.x.x.x
    auto_numbered_title_4 = []  # 形如除去x.和上述3种情况之外的部分

    title_0_format = []  # 形如 " 第X章 XXX" 这种情况
    title_1_format = []
    title_2_format = []
    title_3_format = []
    title_4_format = []

    figure_title_format = []
    table_title_format = []
    xubiaoTitle_format = []

    paragraphs = document.paragraphs
    for each_paragraph in paragraphs:
        text = each_paragraph.text

        auto_numbered_title = get_auto_numbered_title(path)
        for elem in auto_numbered_title:
            priority, auto_numbered_title_content = elem
            if priority == 1:
                auto_numbered_title_1.append(auto_numbered_title_content)
            elif priority == 2:
                auto_numbered_title_2.append(auto_numbered_title_content)
            elif priority == 3:
                auto_numbered_title_3.append(auto_numbered_title_content)
            elif priority == 4:
                auto_numbered_title_4.append(auto_numbered_title_content)
            else:
                pass
        if text.strip()!='':
            res1 = get_tableTitle_format(each_paragraph)
            res2 = get_imageTitle_format(each_paragraph)
            res3 = get_xubiao_format(each_paragraph)
            if len(res1):
                table_title_format.append(res1)
            if len(res2):
                figure_title_format.append(res2)
            if len(res3):
                xubiaoTitle_format.append(res3)

        if len(text) < 30:

            # -------处理自动生成编号------
            if text in auto_numbered_title_1:
                title_1_format.append(get_paragraph_format(each_paragraph))
            elif text in auto_numbered_title_2:
                title_1_format.append(get_paragraph_format(each_paragraph))
            elif text in auto_numbered_title_3:
                title_3_format.append(get_paragraph_format(each_paragraph))
            elif text in auto_numbered_title_4:
                title_4_format.append(get_paragraph_format(each_paragraph))


            # -------处理手动编号------
            elif re.match(_pattern3, text):
                title_3_format.append(get_paragraph_format(each_paragraph))
            elif re.match(_pattern2, text):
                title_2_format.append(get_paragraph_format(each_paragraph))
            elif re.match(_pattern1, text):
                title_1_format.append(get_paragraph_format(each_paragraph))
            elif re.match(_pattern, text):
                title_0_format.append(get_paragraph_format(each_paragraph))
    FORMAT = [title_0_format, title_1_format, title_2_format, title_3_format,
                                title_4_format, table_title_format, figure_title_format, xubiaoTitle_format]
    return FORMAT


def get_paragraph_format(paragraph):
    '''
    通过传入Document.paragraphs中的一个段落，来获得该段落的内容、字体名称、字体大小、段前距离、段后距离、首行缩进、是否居中等信息
    :param paragraph:  Document.paragraphs中的一个段落
    :return: 该段落属性组成的元组
    '''
    """
    暂且在此函数中将Normal样式的默认值做处理(但每次判断段落都要处理，复杂度大。)：
    规则：若函数style.font.name, style.font.Chinese_font_name能获取其默认值，则不做处理，不然用得到文档的默认值赋值处理。
    """

    font_name = []
    font_size = []
    run_text = []
    title_content_and_format = ()
    each_paragraph = paragraph
    text = each_paragraph.text
    style_name = paragraph.style.name

    # 获得继承属性
    paragraph_style_font_name  = get_style_format(paragraph.style)[0]
    paragraph_style_Chinese_font_name= get_style_format(paragraph.style)[1]
    alignment_style = get_style_format(paragraph.style)[2]
    space_before_style = get_style_format(paragraph.style)[3]
    space_after_style = get_style_format(paragraph.style)[4]
    left_indent_style = get_style_format(paragraph.style)[5]
    right_indent_style = get_style_format(paragraph.style)[6]
    first_line_indent_style = get_style_format(paragraph.style)[7]
    # 段落样式的段前距
    space_before = each_paragraph.paragraph_format.space_before
    if space_before == None:
        space_before = space_before_style
    # 段落样式的段后距
    space_after = each_paragraph.paragraph_format.space_after
    if space_after == None:
        space_after = space_after_style
    # 段落样式的对齐方式
    alignment = each_paragraph.paragraph_format.alignment
    if alignment == None:
        alignment = alignment_style
    # 首行缩进
    first_line_indent = each_paragraph.paragraph_format.first_line_indent
    if first_line_indent == None:
        first_line_indent = first_line_indent_style
    # 左缩进 和 右缩进
    left_indent = each_paragraph.paragraph_format.left_indent
    right_indent = each_paragraph.paragraph_format.right_indent
    if left_indent == None:
        first_line_indent = left_indent_style
    if right_indent == None:
        right_indent = right_indent_style

    # 获得每一个run的属性
    """
        获得该段段落样式中的字体、字号等信息，即得到继承样式中的属性,现在知道的是字体样式属性可能会覆盖段落样式的字体属性，则获取字体字号的策略是
    如果该段落的属性值不是None,正常获取。若是None则说明属性是从样式中继承而来，此时先从字体样式中获取，如若为None,则转向字体样式继承的base_style。
    若在没有，转向段落属性，类似的如果段落属性中也没有字体字号信息，转向其base_style，如果在没有，目前还没有找到方法，暂且置为None。
        在python-docx包获得字体名称时font.name 获得的是ascii的值，当样式中规定了中文字体值即eastAsia时，则中文字体的值无法获得。
        策略：获取段落样式中的中文字体名称，若没有，则采用默认的ascii值。
    """
    """"
    if paragraph_style_font_name == None:
        paragraph_style_font_name = each_paragraph.style.base_style.font.name if each_paragraph.style.base_style else None
    if paragraph_style_Chinese_font_name == None:
        paragraph_style_Chinese_font_name = each_paragraph.style.base_style.font.Chinese_font_name if each_paragraph.style.base_style else None
    """
    paragraph_style_font_size = each_paragraph.style.font.size

    if paragraph_style_font_size == None:
        paragraph_style_font_size = each_paragraph.style.base_style.font.size if each_paragraph.style.base_style else None

    for each_run in each_paragraph.runs:
        '''-------size---------'''
        temptext = each_run.text
        run_text.append(temptext)
        run_style_font_size = each_run.style.font.size
        if run_style_font_size == None:
            run_style_font_size = paragraph_style_font_size
        '''---------name------------'''
        each_run_font_name = each_run.font.name
        each_run_Chinese_font_name = each_run.font.Chinese_font_name

        if each_run_font_name == None:
            base_style_font_name = each_run.style.base_style.font.name if each_run.style.base_style else None
            run_style_font_name = each_run.style.font.name if each_run.style.font.name else base_style_font_name
            if run_style_font_name == None:
                run_style_font_name = paragraph_style_font_name

        if each_run_Chinese_font_name == None:
            base_style_Chinese_font_name = each_run.style.base_style.font.Chinese_font_name if each_run.style.base_style else None
            run_style_font_Chinese_font_name = each_run.style.font.Chinese_font_name if each_run.style.font.Chinese_font_name \
                else base_style_Chinese_font_name
            if run_style_font_Chinese_font_name == None:
                run_style_font_Chinese_font_name = paragraph_style_Chinese_font_name


        if each_run.text.isascii():
            '''调试过程中发现的问题，再有的时候中文字体和ascii字符可能会存储在同一个run中，这是按照下列判断会误判'''
            font_name.append(each_run_font_name if each_run_font_name else run_style_font_name)
        else:
            font_name.append(each_run_Chinese_font_name if each_run_Chinese_font_name else run_style_font_Chinese_font_name)

        font_size.append(each_run.font.size if each_run.font.size else run_style_font_size)

    for index, content in enumerate(font_size):
        if content==None:
            pass
        elif content < 100:
            pass
        else:
            font_size[index] = content.pt

    if space_before == None:
        pass
    else:
        space_before = space_before/198120

    if space_after == None:
        pass
    else:
        space_after = space_after/198120

    title_content_and_format = [text,run_text, font_size, font_name, alignment, space_before, space_after, left_indent ,right_indent,first_line_indent]

    return title_content_and_format

def get_none_inherent_fornmat(paragraph):
    '''
       通过传入Document.paragraphs中的一个段落，来获得该段落的内容、字体名称、字体大小、段前距离、段后距离、首行缩进、是否居中等信息
       :param paragraph:  Document.paragraphs中的一个段落
       :return: 该段落属性组成的元组
       '''
    font_name = []
    font_size = []
    title_content_and_format = ()
    each_paragraph = paragraph
    text = each_paragraph.text
    alignment = each_paragraph.alignment
    if each_paragraph.paragraph_format.space_before == None:
        space_before = None
    else:
        space_before = each_paragraph.paragraph_format.space_before / 198120
    if each_paragraph.paragraph_format.space_after == None:
        space_after = None
    else:
        space_after = each_paragraph.paragraph_format.space_after / 198120
    # 获得每一个run的属性
    #paragraph_style_font_size = each_paragraph.style.font.size

    for each_run in each_paragraph.runs:
        #run_style_font_size = each_run.font.size
        #run_style_font_name = each_run.font.name
        font_name.append(each_run.font.name)
        '''
        if each_run.font.size == None:
            font_size.append(None)
        else:
            font_size.append(each_run.font.size.pt)
            '''
        font_size.append(each_run.font.size)
    if each_paragraph.paragraph_format.first_line_indent == None:
        first_line_indent = None
    else:
        first_line_indent = each_paragraph.paragraph_format.first_line_indent
    title_content_and_format = [text, font_size, font_name, space_before, space_after, first_line_indent, alignment]

    return title_content_and_format


def get_auto_numbered_title(path):
    '''
    说明：自动编号的标题段落属性中<w:pStyle w:val="ListParagraph"/>
    一级标题：<w:numPr>
                <w:ilvl w:val="0"/> # 1.
                <w:ilvl w:val="1"/> # 1.1
                <w:ilvl w:val="2"/> # 1.1.1
                <w:ilvl w:val="3"/> # 1.1.1.1
                        ...
            </w:numPr>
    要想获取自动编号的标题，先判断段落属性，接着进行级别的判断，再提取文本，最后与非自动编号的标题合并。

    :param path: docx的文件绝对路径
    :return: 一个元组content里面包含标题的级别和名称,eg (1，'这是一个二级标题1.1'）
    '''
    document = Document(path)
    content =[]
    for child in document.element.body.iter():
        #print(child.tag)
        if child.tag.endswith('main}p'):
            #print(child.text)
            _text =''
            for elem in child.iter():
                if elem.tag.endswith('t'):
                    if elem.text == None:
                        pass
                    else:
                        _text += elem.text
                    #print(_text)
            for elem in child.iter():
                if elem.tag.endswith('pPr'):
                    if hasattr(elem,'pStyle'):
                        if elem.pStyle ==None:
                            pass
                        elif elem.pStyle.val=='ListParagraph':
                            if elem.numPr != None:
                                Priority = elem.numPr.ilvl.val
                                content.append((Priority,_text))
    return content


    # return content


def get_styles(path):
    '''
    获得该文档的样式styles
    :param path: docx文档的绝对路径
    :return: 样式列表
    '''
    document = Document(path)
    s = []
    for paragraph in document.paragraphs:
        font_name = paragraph.style.font.name
        if font_name == None:

            font_name = paragraph.style.base_style.name if paragraph.style.base_style else 'no base_style, self has no fontname'
        content = (paragraph.text, paragraph.style.name,font_name)
        s.append(content)
    return s


def get_cover_fmt(path):
    '''
    由于封面页数比较靠前，直接在前20段做判断，避免过多的时间消耗
    :param path: docx文件的绝对路径
    :return: 属性列表 content_and_fmt
                内容提要页的段索引 end_index
    '''
    document = Document(path)
    paragraphs = document.paragraphs
    # end_index 是"内容提要"关键字的段索引
    end_index = 0
    content_and_fmt = []

    # 教材编号
    textbook_SN_pattern = r'\w\-\w\-\w{5}\-\w{2}'
    for index, paragraph in enumerate(paragraphs):
        if index < 130:
            title =  paragraph.text.replace(' ','')
            if re.match(textbook_SN_pattern,title):
                start_index = index
                break
            else:
                return None
        else:
            return None
    for index, paragraph in enumerate(paragraphs):
        if index < 130:
            if paragraph.text.replace(' ','') == '内容提要':
                end_index = index
                break
        else:
            return None
    for index, paragraph in enumerate(paragraphs):
        if index in range(start_index, end_index) and paragraph.text.replace(' ', '') != '':
            content_and_fmt.append(get_paragraph_format(paragraph))
    return content_and_fmt

def get_informative_abstract_fmt(path):
    '''

    :param path: docx文档的绝对路径
    :return: 内容提要页属性
    '''
    document = Document(path)
    paragraphs = document.paragraphs
    content_and_fmt = []
    next_part_title = '教材体系工程'
    for index, paragraph in enumerate(paragraphs):
        if index < 130:
            if paragraph.text.replace(' ','') == '内容提要':
                start_index = index
                break
        else:
            return None
    for index, paragraph in enumerate(paragraphs):
        if index < start_index:
            continue
        if index < 160:
            if next_part_title in paragraph.text.replace(' ',''):
                end_index = index
                break
        else:
            return None
    for index, paragraph in enumerate(paragraphs):
        if index in range(start_index, end_index) and paragraph.text.replace(' ', '') != '':
            content_and_fmt.append(get_paragraph_format(paragraph))
    return content_and_fmt

def get_BSZ_fmt(path):
    '''
    BSZ:编审组缩写
    :param path: docx文件绝对路径
    :return: 编审组页的段属性
    '''
    document = Document(path)
    paragraphs = document.paragraphs
    content_and_fmt = []
    next_part_title = '前言'
    '''
    for index, paragraph in enumerate(paragraphs):
        if index < 100 :
            text_temp = paragraph.text.replace(' ','')
            if '印数' in text_temp:
                start_index1 = index
                break
        else:
            return None
    '''
    for index, paragraph in enumerate(paragraphs):

        if index < 200 :
            text_temp = paragraph.text.replace(' ','')
            if '编审组' in text_temp:
                start_index = index-1
                break
        else:
            return None
    for index, paragraph in enumerate(paragraphs):
        if index < start_index:
            continue
        if index < 250:
            if next_part_title in paragraph.text.replace(' ',''):
                end_index = index
                break
        else:
            return None
    for index, paragraph in enumerate(paragraphs):
        if index in range(start_index, end_index) and paragraph.text.replace(' ', '') != '':
            content_and_fmt.append(get_paragraph_format(paragraph))
    return content_and_fmt

def get_preface_fmt(path, start_index=0):
    '''
    BSZ:编审组缩写
    :param path: docx文件绝对路径
    :return: 前言页的段属性
    '''
    document = Document(path)
    paragraphs = document.paragraphs
    content_and_fmt = []
    next_part_title = '目录'
    # 获取该页的起始页
    for index, paragraph in enumerate(paragraphs):
        if index < 200:
            if paragraph.text.replace(' ','') == '前言':
                start_index = index
                break
        else:
            return None
    for index, paragraph in enumerate(paragraphs):
        if index < start_index:
            continue
        if index < 250:
            if next_part_title == paragraph.text.replace(' ',''):
                end_index = index
                break
        else:
            return None
    for index, paragraph in enumerate(paragraphs):
        if index in range(start_index, end_index-1) and paragraph.text.replace(' ', '') != '':
            content_and_fmt.append(get_paragraph_format(paragraph))
    return content_and_fmt


def get_SIKAOTI_fmt(path):
    '''
    得到思考题的格式，策略是当遇到"复习思考题"字眼时
    :param path文件路径
    :return  思考题格式列表
    '''
    doc = Document(path)
    paragraphs = doc.paragraphs
    pattern = r'^\d+'
    SIKAOTI_index_list = []
    SIKAOTI_title_fmt = []
    SIKAOTI_content_fmt=[]
    for index, each_paragraph in enumerate(paragraphs):
        if each_paragraph.text == '复习思考题':
            SIKAOTI_index_list.append(index)
            SIKAOTI_title_fmt.append(get_paragraph_format(each_paragraph))
            # 思考题内容编号 list_content_index
            list_content_index = index+1
            while(re.match(pattern,paragraphs[list_content_index].text)):
                SIKAOTI_content_fmt.append(get_paragraph_format(paragraphs[list_content_index]))
                SIKAOTI_index_list.append(list_content_index)
                list_content_index = list_content_index +1

    return [SIKAOTI_title_fmt,SIKAOTI_content_fmt,SIKAOTI_index_list]


def get_title_index(path):
    '''
        title_list:各级标题和图表标题的文字内容
        title_index_list：各级标题和图表标题的索引
        body_index：目前是从0开始去除 title_index_list 的正文索引
    '''
    doc = Document(path)
    pars = doc.paragraphs
    title_list = []
    fmt = get_numbered_title(path)
    title_index_list = []
    body_index = []
    for i in range(len(fmt)):
            for j in range(len(fmt[i])):
                title_list.append(fmt[i][j][0])

    for m in range(len(pars)):
        body_index.append(m)
        if pars[m].text in title_list:
            title_index_list.append(m)
    for n in range(len(title_index_list)):
        del body_index[title_index_list[n]]
    return title_index_list


def process_main_body(path,standard,title_index):
    '''
    一般情况下，目录和附录之间的内容为正文内容，正文内容中包含有表格等，先获得正文的段落范围,
    判断正文时要将正文中的标题（章节标题、节标题、款标题、图表标题、包括章后思考题等)去除,单独判断这些标题的样式，
    加入单独的错误列表,获得标题的段落索引，当遇到这些索引时，跳过即可。

    :param: path
    :param: standard
    retrun
    '''
    # --------标准--------
    '''
    font_size_std = standard[0]
    font_name_std = standard[1]
    alignment_std = standard[2]
    para_before = standard[3]
    para_after = standard[4]
    left_indentation = standard[5]
    right_indentation = standard[6]
    '''
    # ---------return内容-----------
    all_errors = []
    # ==========================
    start_index = -1
    end_index = -1
    paras = Document(path).paragraphs
    standard = standard.split()
    '''获取正文的段索引范围'''
    for index,para in enumerate(paras):
        if re.match(r'\第\.\章',para.text):
            start_index= index
        if para.text.replace(' ','') == '附录':
            end_index = index
    '''
    当无法判断其前后索引时，认为此文章整体均为正文内容，因为其判定方法为|第X章|和|附录|为界定点，
    所以这两个部分内容必须存在！！！！
    '''
    if start_index == -1:
        start_index = 0
    if end_index == -1:
        end_index = len(Document(path).paragraphs)-1
    for para_index,para in enumerate(Document(path).paragraphs):
        if para_index < start_index:
            continue
        elif para_index > end_index:
            break
        elif para_index in title_index:
            continue
        else:
            fmt = get_paragraph_format(para)
            error = error_process_unit_for_each_para(fmt,standard,'TEXT')
            all_errors.append(error)
    return all_errors


def process_title(fmt,standard,type):
    '''
    处理一级标题，并报告错误，包括连续性错误？
    :param fmt:包含级标题的格式列表
    :param standard:
    :param type:title的类型，是一级标题还是二级标题等或者说是图表标题...
    :return: 错误信息列表[['text',[字体大小错误run],[字体名称错误run],bool,bool,bool,bool,bool...],['text',[字体大小错误run],[字体名称错误run],bool,bool,bool,bool,bool...],...]
    '''
    all_errors = []

    for i in fmt:
        res = error_process_unit_for_each_para(i,standard,type)
        all_errors.append(res)


    return all_errors


def error_process_unit(category,detail,paragraph_fmt_elem, standard_config):
    '''

    :param category: 文档的内容类别
    :param fmt: 输入的该段的段落格式列表
    :param standard: 该类型内容下的每个小种类的标准
    :return: 错误列表
    '''
    standard = standard_config[category][detail].split()
    standard[2] = float(standard[2])
    standard[3] = float(standard[3])
    standard[4] = float(standard[4])
    standard[5] = float(standard[5])
    standard[6] = float(standard[6])


    font_size_error =[]
    font_name_error=[]
    alignment_error = False
    format_font_size_pt = {'八号': 5, '七号': 5.5, '小六': 6.5, '六号': 7.5, '小五': 9, '五号': 10.5,
                           '小四': 12, '四号': 14, '小三': 15, '三号': 16, '小二': 18, '二号': 22, '一号': 26, '小初': 36,'初号':42}
    # 判断字号
    '''之加入错误的内容，不显示具体错误信息，以段为单位进行报错'''
    font_size = format_font_size_pt[standard[0]]
    font_name = standard[1]
    for index0, elem0 in enumerate(paragraph_fmt_elem[2]):
        if elem0 != font_size:
            font_size_error.append(paragraph_fmt_elem[1][index0])
        # 判断字体名称
    for index1,elem1 in enumerate(paragraph_fmt_elem[3]):
        if elem1 != font_name:
            font_name_error.append(paragraph_fmt_elem[1][index1])
            # 判断居中方式
    alignment = paragraph_fmt_elem[4] if paragraph_fmt_elem[4] else 0
    if alignment != standard[2]:
        alignment_error = True
    return [paragraph_fmt_elem[0],font_size_error, font_name_error,alignment_error]


def error_process_unit_for_each_para(paragraph_fmt_elem, standard,category):
    '''

    :param paragraph_fmt_elem: 具体需要判断的某一项，如：AUTHOR中的编审组[处理每一段的内容]
    :param standard: 针对这一具体内容的标准,其格式为列表['四号','黑体','0','0','0','0'，0'](用split函数即可）
    :param category: 用作提示错误输出的类别，例如'AUTHOR'
    :return: 错误列表
    '''
    standard[2] = float(standard[2])
    standard[3] = float(standard[3])
    standard[4] = float(standard[4])
    standard[5] = float(standard[5])
    standard[6] = float(standard[6])

    # 错误列表
    font_size_error = []
    font_name_error = []
    alignment_error = False
    space_before_error = False
    space_after_error = False
    left_indent_error = False
    right_indent_error = False
    first_line_indent_error = False

    format_font_size_pt = {'八号': 5, '七号': 5.5, '小六': 6.5, '六号': 7.5, '小五': 9, '五号': 10.5,
                           '小四': 12, '四号': 14, '小三': 15, '三号': 16, '小二': 18, '二号': 22, '一号': 26, '小初': 36}
    # 判断字号
    font_size = format_font_size_pt[standard[0]]
    font_name = standard[1]

    # paragraph_fmt_elem[0]为该段的内容
    # paragraph_fmt_elem[1]为每个run的内容
    # paragraph_fmt_elem[2]为每个run的字号
    # paragraph_fmt_elem[3]为每个run的字体名称

    for index,elem in enumerate(paragraph_fmt_elem[2]):
        if elem != font_size:
            font_size_error.append(paragraph_fmt_elem[1][index])
    # 判断字体名称
    for index1,elem1 in enumerate(paragraph_fmt_elem[3]):
        if elem1 != font_name:
            font_name_error.append(paragraph_fmt_elem[1][index1])
    # 判断居中方式
    alignment = paragraph_fmt_elem[4] if paragraph_fmt_elem[4] else 0
    if alignment != standard[2]:
        alignment_error = True
    # 判断段前距
    space_before = paragraph_fmt_elem[5] if paragraph_fmt_elem[5] else 0
    if space_before != standard[3]:
        space_before_error = True
    # 判断段后距
    space_after = paragraph_fmt_elem[6] if paragraph_fmt_elem[7] else 0
    if space_after != standard[4]:
        space_after_error = True
    # 判断左缩进
    left_indent = paragraph_fmt_elem[7] if paragraph_fmt_elem[7] else 0
    if left_indent != standard[5]:
        left_indent_error = True
    # 判断右缩进
    right_indent = paragraph_fmt_elem[8] if paragraph_fmt_elem[8] else 0
    if right_indent != standard[6]:
        right_indent_error = True
    # 判断首行缩进
    '''
    first_line_indent = paragraph_fmt_elem[9] if paragraph_fmt_elem[9] else 0
    if first_line_indent != standard[7]:
        error = category + ': ' + paragraph_fmt_elem[0] + ' 首行错误，应为' + \
                str(standard[7])
        first_line_indent_error.append(error)
        '''

    return [paragraph_fmt_elem[0],font_size_error, font_name_error, alignment_error,space_before_error,space_after_error,left_indent_error\
            ,right_indent_error,first_line_indent_error]


def error_process_unit_two_parts(part1,fmt,category,standard1,standard2):
    '''
    处理一个自然段中，段首字符与剩余内容格式要求不同的情况。例如编审组页中组长、副组长、组员标题的判断。
    :paras part1 去除空格后的第一部分内容
    :paras part2 去除空格后的第二部分内容
    :fmt 该段落的格式列表
    :paras category 所属的类别
    :paras standard1 第一部分标准
    :paras standard2 第二部分标准
    :return  error_list 错误列表
    '''
    font_name_error = []
    font_size_error = []
    standard1[2] = float(standard1[2])
    standard1[3] = float(standard1[3])
    standard1[4] = float(standard1[4])
    standard1[5] = float(standard1[5])
    standard1[6] = float(standard1[6])
    standard2[2] = float(standard2[2])
    standard2[3] = float(standard2[3])
    standard2[4] = float(standard2[4])
    standard2[5] = float(standard2[5])
    standard2[6] = float(standard2[6])
    format_font_size_pt = {'八号': 5, '七号': 5.5, '小六': 6.5, '六号': 7.5, '小五': 9, '五号': 10.5,
                           '小四': 12, '四号': 14, '小三': 15, '三号': 16, '小二': 18, '二号': 22, '一号': 26, '小初': 36}
    temp = ''
    check_index = [-1]
    for index_part1, content_part1 in enumerate(fmt[1]):
        temp += content_part1.replace(' ','')
        if temp == part1:
            check_index[0] = index_part1
            break
        # ==========================处理part1===============================
        # 判断字体大小
        # content[2]是每个run中字体大小的集合
    for index_font_size, run_font_size in enumerate(fmt[2]):
            if index_font_size <= check_index[0]:
                if run_font_size != format_font_size_pt[standard1[0]]:
                    font_size_error.append(category+part1+':'+ '"'+fmt[1][index_font_size]+'"'+'，字体大小不正确，应为' + standard1[0])
    # 判断字体名称
    # content[3]是每个run中字体名称的集合
    for index_font_name, run_font_name in enumerate(fmt[3]):
            if index_font_name <= check_index[0]:

                if run_font_name != standard2[1]:
                    font_name_error.append(category+':'+ '"'+fmt[1][index_font_name]+'"'+'字体不正确，应为' + standard2[1])
    # ==========================处理part2===============================
    # 判断字体大小
    # content[2]是每个run中字体大小的集合
    for index_font_size, run_font_size in enumerate(fmt[2]):
        if index_font_size > check_index[0]:
            if run_font_size != format_font_size_pt[standard2[0]]:
                font_size_error.append(category+':'+ '"'+fmt[1][index_font_size]+'"'+' 字体大小不正确，应为' + standard2[0])
    # 判断字体名称
    # content[3]是每个run中字体名称的集合
    for index_font_name, run_font_name in enumerate(fmt[3]):
        if index_font_name > check_index[0]:
            if run_font_name != standard2[1]:
                font_name_error.append(category+':'+ '"'+fmt[1][index_font_name]+'"'+'字体不正确，应为' + standard2[1])
    # 判断居中方式
    # 段前距
    # 段后距
    # 左缩进
    # 右缩进
    # 首行缩进
    return [font_size_error,font_name_error]


def process_preface(fmt,standard,type):
    '''
    处理前言格式模块，输出的是前言的格式错误。前言分四个检测的部分：前言标题，正文，编者，时间。
    :params fmt 前言的格式列表
    :params standard 前言标准
    :params type 前言中需要判断的四个类型
    '''
    if fmt == None:
        return []
    res = []
    preface_standard_title = standard['PREFACE']['前言'].split()
    preface_standard_main_part = standard['PREFACE']['正文'].split()
    preface_standard_author = standard['PREFACE']['编者'].split()
    preface_standard_date = standard['PREFACE']['时间'].split()
    date_pattern = r'^\d{4}\年\d+\月'
    for each_paragraph in fmt:
        if each_paragraph[0].replace(' ','') =='前言':
            error0 = error_process_unit_for_each_para(each_paragraph,preface_standard_title,'前言标题')
            res.append(error0)
        elif each_paragraph[0].startswith('编者'):
            error1 = error_process_unit_for_each_para(each_paragraph,preface_standard_author,'前言作者')
            res.append(error1)
        elif re.match(date_pattern,each_paragraph[0]):
            error2 = error_process_unit_for_each_para(each_paragraph, preface_standard_date, '前言日期')
            res.append(error2)
        else:
            error3 = error_process_unit_for_each_para(each_paragraph, preface_standard_main_part, '前言正文')
            res.append(error3)

    return res


def title_continuity():
    '''
    检测章节标题一致性
    '''
    pass



def table_and_figure_continuity():
    '''
    检测图标题、表标题一致性
    '''
    pass


def consistency(path):
    '''
    检测书名主编日期一致性
    '''
    pass


def title_continuity(path):
    '''
        title_text_list :去掉图表标题的文字列表
        title_index_list：各级标题和图表标题的索引
        unprocessed_title_list:未处理的正序、不含图表标题 的文字列表
        title_list: 处理后 正序、去掉汉字等字符、只剩数字和点 的标题列表

        detect_result： 形式：[[缺少的标题],[[顺序错误的标题]]]，故分别取为detect_result[0] 与 detect_result[1][0]
      '''
    doc = Document(path)
    pars = doc.paragraphs
    #----------得到所有标题的格式--------
    all_fmt = get_numbered_title(path)
    #----------去掉图表标题----------
    fmt = all_fmt[0:4]
    '''----------得到正序的标题列表-----------'''
    title_text_list = []
    title_index_list = []
    unprocessed_title_list = []
    title_list =[]
    for i in range(len(fmt)):
        for j in range(len(fmt[i])):
            title_text_list.append(fmt[i][j][0])
     #------------判断如果标题文字在某段中，则获得索引，这样可保证列表中标题的正序--------------
    for m in range(len(pars)):
        if pars[m].text in title_text_list:
            title_index_list.append(m)
    #-------利用正序索引列表得到正序标题文字列表-------------
    for n in range(len(title_index_list)):
        unprocessed_title_list.append(pars[title_index_list[n]].text)
    '''-----对标题列表进行处理，去掉点外的各种符号，并去掉一级标题中的“第”字，使其只含数字和点------------'''
    r = '[’!"#$%&\'（）*+,/:;<=>?@[\\]^\-_`{|}~\n。！， ]+'
    for j in range(len(unprocessed_title_list)):
        unprocessed_title_list[j] = re.sub(r, '', unprocessed_title_list[j])
        if re.match(r'^第\w{1,2}?章', unprocessed_title_list[j]):
            unprocessed_title_list[j] = unprocessed_title_list[j][1:]  # 去掉“第”字
    #print(unprocessed_title_list)
    for i in range(len(unprocessed_title_list)):
        title_list.append(extract_nums(unprocessed_title_list[i])) #得到纯数字和点的标题列表
    #print(title_list)

    ''' ----进行判断-----用到的函数有----extract_nums----title_continuous-------point_to_number-----output_miss_list----'''
    detect_result = title_continuous(title_list)
    # print(detect_result)
    #print('缺少的标题有：', detect_result[0])
    #print('顺序错误的标题为：', detect_result[1][0])
    return detect_result, all_fmt




def extract_nums(str_text):
    #-------去掉标题文字--------
    for i,t in enumerate(str_text):
        # if t.isdigit():
        if t in '0123456789.':
            continue
        else:
            return str_text[:i]
    return str_text


def title_continuous(catalogue_number_list):
    for j in range(len(catalogue_number_list)):
        if catalogue_number_list:
            if catalogue_number_list[j]:
                if catalogue_number_list[j][-1] =='.':
                    catalogue_number_list[j] = catalogue_number_list[j][:-1]
                elif catalogue_number_list[j][0] == '.':
                    catalogue_number_list[j] = catalogue_number_list[j][1:]
    temp_list = catalogue_number_list[:]
    for i in temp_list:
        if i:
            for j in i:
                if not j.isdigit() and j != '.':
                    catalogue_number_list.remove(i)
                    break
                elif i[0] == '0':
                    catalogue_number_list.remove(i)
                    break
                else:
                    pass
        else:
            catalogue_number_list.remove(i)
    list_miss = []
    list_discrete = []
    if catalogue_number_list:
        list_che = []
        if catalogue_number_list[0] != '1':
            list_miss.append('1')
        for i in catalogue_number_list:
            if not list_che:
                list_che.append(i)
                continue
            if i.count('.') == list_che[-1].count('.'):
                num_i = point_to_number(i)
                num_list = point_to_number(list_che[-1])
                if len(str(num_i)) != len(str(num_list)):
                    m = abs(len(str(num_i)) - len(str(num_list)))
                    if len(str(num_i)) < len(str(num_list)):
                        for j in range(m):
                            num_i *= 10
                    else:
                        for j in range(m):
                            num_list *= 10
                if num_i - num_list == 1:
                    list_che.pop()
                    list_che.append(i)
                elif num_i - num_list > 1:
                    for kk in output_miss_list(list_che[-1], i):
                        if kk not in list_miss:
                            list_miss.append(kk)
                    list_che.pop()
                    list_che.append(i)
                elif num_i - num_list < 1:
                    temp = []
                    temp.append(list_che[-1])
                    temp.append(i)
                    list_discrete.append(temp)
                    list_che.pop()
                    list_che.append(i)
            elif i.count('.') > list_che[-1].count('.'):
                num_i = point_to_number(i)
                num_list = point_to_number(list_che[-1])
                if len(str(num_i)) != len(str(num_list)):
                    m = abs(len(str(num_i)) - len(str(num_list)))
                    if len(str(num_i)) < len(str(num_list)):
                        for j in range(m):
                            num_i *= 10
                    else:
                        for j in range(m):
                            num_list *= 10
                if num_i - num_list == 1:
                    # list_che.pop()
                    list_che.append(i)
                elif num_i - num_list > 1:
                    for kk in output_miss_list(list_che[-1], i):
                        if kk not in list_miss:
                            list_miss.append(kk)
                            # mid = []
                            # mid.append(int(dic[list_che[-1]]))
                            # mid.append(int(dic[i]))
                            # list_miss_page.append(mid)
                    # list_che.pop()
                    list_che.append(i)
                elif num_i - num_list < 1:
                    temp = []
                    temp.append(list_che[-1])
                    temp.append(i)
                    list_discrete.append(temp)
                    list_che.append(i)
            elif i.count('.') < list_che[-1].count('.'):
                while i.count('.') != list_che[-1].count('.'):
                    list_che.pop()
                    if not list_che:
                        list_che.append(i)
                        break
                if point_to_number(i) - point_to_number(list_che[-1]) == 1:
                    list_che.pop()
                    list_che.append(i)
                elif point_to_number(i) - point_to_number(list_che[-1]) > 1:
                    for kk in output_miss_list(list_che[-1], i):
                        if kk not in list_miss:
                            list_miss.append(kk)
                    list_che.pop()
                    list_che.append(i)
                elif point_to_number(i) - point_to_number(list_che[-1]) < 1:
                    temp = []
                    temp.append(list_che[-1])
                    temp.append(i)
                    list_discrete.append(temp)
                    list_che.pop()
                    list_che.append(i)
        return list_miss, list_discrete
    else:
        return [],[]
def point_to_number(s):
    res = ''
    for i in s:
        if i.isdigit():
            res += i
    if res:
        return int(res)
    else:
        return False
def output_miss_list(a, b):
    temp_a = a.split('.')
    temp_b = b.split('.')
    min_length = min(len(temp_a), len(temp_b))
    for i in range(min_length):
        if int(temp_a[i]) > int(temp_b[i]):
            return []
    temp_str = ''
    res = []
    if len(temp_a) != len(temp_b):
        n = max(len(temp_a),len(temp_b)) - min(len(temp_a),len(temp_b))
        if len(temp_a) < len(temp_b):
            for i in range(n):
                temp_a.append('0')
        else:
            for i in range(n):
                temp_b.append('0')
    for i in range(len(temp_a)):
        if temp_a[i] != temp_b[i]:
            m = i
            for k in range(m+1 ,len(temp_a)):
                temp_a[k] = '0'
            break

    while temp_str != b:
        temp_str = ''
        for i in range(len(temp_a)):
            if temp_a[i] == temp_b[i]:
                temp_str += temp_a[i]
                temp_str += '.'
            else:
                temp_str += str(int(temp_a[i]) + 1)
                temp_a[i] = str(int(temp_a[i]) + 1)
                if temp_str != b:

                    res.append(temp_str)
                break
        if temp_str[-1] == '.':
            temp_str = temp_str[:-1]
    return res


def decide_style_normal(path):
    '''

    '''
    doc = Document(path)
    normal = doc.styles['Normal']
    if normal.font.name == None:
        normal.font.name = get_doc_default_Pr.font_name(path)
        global normal_font_name
        normal_font_name = normal.font.name
    if normal.font.Chinese_font_name == None:
        normal.font.Chinese_font_name = get_doc_default_Pr.Chinese_font_name(path)
        global normal_Chinese_font_name
        normal_Chinese_font_name = normal.font.Chinese_font_name
    return normal_font_name, normal_Chinese_font_name

if __name__ =='__main__':

    path = r'H:\DocFormatCheckProject\Word-check\move\版式规范\1机务系统教材版式规范模板-Word版-20190719.docx'
    #path = r'G:\v1\files_for_testing\无兼容性2机务系统教材版式规范模板-Word版-20190719.docx'
    doc = Document(path)
    normal = doc.styles['Normal']
    if normal.font.name == None:
        normal.font.name = get_doc_default_Pr.font_name(path)
        normal_font_name = normal.font.name
    if normal.font.Chinese_font_name == None:
        normal.font.Chinese_font_name = get_doc_default_Pr.Chinese_font_name(path)
        normal_Chinese_font_name = normal.font.Chinese_font_name
    print(normal.font.name,normal.font.Chinese_font_name)
    #RES = title_continuity(path)
    #print(RES)
    show_all_content_fmt(path)









